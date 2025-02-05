from django.shortcuts import render, redirect, get_object_or_404
import requests
from django.shortcuts import render
from .forms import AddressForm
from .models import *
from django.http import JsonResponse
from .forms import *
from rest_framework import viewsets
from .serializers import *
from django.http import HttpResponse
from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.serializers import serialize
import requests
from decimal import Decimal
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import docx
from io import BytesIO
from functools import wraps
from django.http import HttpResponseForbidden

def role_required(roles):
    def decorator(view_func):
        @wraps(view_func)
        def _wrapped_view(request, *args, **kwargs):
            if request.user.is_authenticated and request.user.idRol and request.user.idRol.nombreRol in roles:
                return view_func(request, *args, **kwargs)
            return HttpResponseForbidden("No tienes permiso para acceder a esta página")
        return _wrapped_view
    return decorator

#SERIALIZERS (API):
class RolUsuarioViewSet(viewsets.ModelViewSet):
    queryset = rolUsuario.objects.all()
    serializer_class = RolUsuarioSerializer

class RegionViewSet(viewsets.ModelViewSet):
    queryset = Region.objects.all()
    serializer_class = RegionSerializer

class ComunaViewSet(viewsets.ModelViewSet):
    queryset = Comuna.objects.all()
    serializer_class = ComunaSerializer

class UsuarioCustomViewSet(viewsets.ModelViewSet):
    queryset = UsuarioCustom.objects.all()
    serializer_class = UsuarioCustomSerializer

class TallerViewSet(viewsets.ModelViewSet):
    queryset = Taller.objects.all()
    serializer_class = TallerSerializer

class MarcaViewSet(viewsets.ModelViewSet):
    queryset = Marca.objects.all()
    serializer_class = MarcaSerializer

class TipoVehiculoViewSet(viewsets.ModelViewSet):
    queryset = TipoVehiculo.objects.all()
    serializer_class = TipoVehiculoSerializer

class VehiculoViewSet(viewsets.ModelViewSet):
    queryset = Vehiculo.objects.all()
    serializer_class = VehiculoSerializer

class TipoAgendaViewSet(viewsets.ModelViewSet):
    queryset = TipoAgenda.objects.all()
    serializer_class = TipoAgendaSerializer

class AgendaViewSet(viewsets.ModelViewSet):
    queryset = Agenda.objects.all()
    serializer_class = AgendaSerializer

def index(request):
	return render(request, 'core/index.html')

@role_required(["Administrador"])
def administrar_talleres(request):
    talleres = Taller.objects.all()
    return render(request, 'core/administrar_talleres.html', {'talleres': talleres})

@role_required(["Administrador"])
def crear_taller(request):
    form = TallerForm()
    coordinates = {}
    address = ""

    if request.method == 'POST':
        if 'crear_taller' in request.POST:
            form = TallerForm(request.POST, request.FILES)
            if form.is_valid():
                taller = form.save(commit=False)
                latitud = request.POST.get('latitud')
                longitud = request.POST.get('longitud')
                direccion = request.POST.get('direccion')
                if latitud and longitud:
                    taller.latitud = float(latitud)
                    taller.longitud = float(longitud)
                if direccion:
                    taller.direccion = direccion
                taller.save()
                return redirect('administrar_talleres')
        elif 'obtener_ubicacion' in request.POST:
            address = request.POST.get('address')
            api_key = '855d3348e0ee4298b17b00e020377da5'
            base_url = 'https://api.opencagedata.com/geocode/v1/json'
            params = {'q': address, 'key': api_key}
            response = requests.get(base_url, params=params)
            if response.status_code == 200:
                data = response.json()
                if data['results']:
                    location = data['results'][0]['geometry']
                    address_details = data['results'][0]['components']
                    if 'state' in address_details and address_details['state'] == 'Santiago Metropolitan Region':
                        coordinates = {
                            'lat': location['lat'],
                            'lng': location['lng']
                        }
                        return JsonResponse({'success': True, 'coordinates': coordinates, 'address': address})
                    else:
                        return JsonResponse({'success': False, 'error': 'La dirección seleccionada no está en la Región Metropolitana.'})
            return JsonResponse({'success': False, 'error': 'Error en la solicitud de geocodificación.'})

    return render(request, 'core/crear_taller.html', {'form': form, 'coordinates': coordinates, 'address': address})

@role_required(["Administrador"])
def modificar_taller(request, id_taller):
    taller = get_object_or_404(Taller, idTaller=id_taller)
    if request.method == 'POST':
        if 'modificar_taller' in request.POST:
            form = TallerForm(request.POST, request.FILES, instance=taller)
            if form.is_valid():
                taller_modificado = form.save(commit=False)
                # Aquí actualizamos la dirección si se recibió una nueva desde el formulario
                if 'direccion' in request.POST:
                    taller_modificado.direccion = request.POST['direccion']
                taller_modificado.save()
                return redirect('administrar_talleres')
        elif 'obtener_ubicacion' in request.POST:
            address = request.POST.get('address')
            api_key = '855d3348e0ee4298b17b00e020377da5'
            base_url = 'https://api.opencagedata.com/geocode/v1/json'
            params = {'q': address, 'key': api_key}
            response = requests.get(base_url, params=params)
            if response.status_code == 200:
                data = response.json()
                if data['results']:
                    location = data['results'][0]['geometry']
                    address_details = data['results'][0]['components']
                    if 'state' in address_details and address_details['state'] == 'Santiago Metropolitan Region':
                        coordinates = {
                            'lat': location['lat'],
                            'lng': location['lng']
                        }
                        # Actualizamos la dirección en el objeto taller y lo guardamos
                        taller.direccion = address
                        taller.save()
                        return JsonResponse({'success': True, 'coordinates': coordinates, 'address': address})
                    else:
                        return JsonResponse({'success': False, 'error': 'La dirección seleccionada no está en la Región Metropolitana.'})
            return JsonResponse({'success': False, 'error': 'Error en la solicitud de geocodificación.'})

    else:
        form = TallerForm(instance=taller)
    
    return render(request, 'core/modificar_taller.html', {'form': form, 'taller': taller})

@role_required(["Administrador"])
def eliminar_taller(request, id_taller):
    taller = get_object_or_404(Taller, idTaller=id_taller)
    taller.delete()
    return redirect('administrar_talleres')

@login_required
def agendar_hora(request):
	return render(request, 'core/agendar_hora.html')

@login_required
def agendar(request):
	return render(request, 'core/agendar.html')

@login_required
def contactanos(request):
    if request.method == 'POST':
        form = ContactoForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Mensaje enviado correctamente.')
            return redirect('contactanos')
        else:
            messages.error(request, 'Hay algún problema con los datos ingresados, revise nuevamente.')
    else:
        form = ContactoForm()
    return render(request, 'core/contactanos.html', {'form': form})

def login(request):
	return render(request, 'core/login.html')

@login_required
def mapa(request):
    talleres = Taller.objects.all()
    talleres_json = serialize('json', talleres, fields=('nombreTaller', 'direccion', 'latitud', 'longitud'))
    return render(request, 'core/mapa.html', {'talleres_json': talleres_json})

@login_required
def mis_reservas(request):
    agendas = Agenda.objects.filter(cliente=request.user)
    return render(request, 'core/mis_reservas.html', {'agendas': agendas})

@csrf_exempt
def actualizar_estado_agenda(request, id_agenda):
    if request.method == 'POST':
        agenda = get_object_or_404(Agenda, idAgenda=id_agenda)
        
        # Obtener la instancia de EstadoAgenda correspondiente
        try:
            estado_pagado = EstadoAgenda.objects.get(nombreEstado='Pagado')
        except EstadoAgenda.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Estado "Pagado" no encontrado'}, status=400)

        agenda.estado = estado_pagado
        agenda.save()

        # Devolver una respuesta JSON con success
        return JsonResponse({'success': True, 'filename': f'Reserva_{agenda.idAgenda}.docx'})

    return JsonResponse({'success': False}, status=400)

def generar_documento_word(request, id_agenda):
    agenda = get_object_or_404(Agenda, idAgenda=id_agenda)

    # Generar el documento Word
    doc = docx.Document()
    doc.add_heading('Detalle de la Reserva', level=1)
    doc.add_paragraph(f'Fecha de Atención: {agenda.fechaAtencion}')
    doc.add_paragraph(f'Hora de Atención: {agenda.horaAtencion}')
    doc.add_paragraph(f'Tipo de Agenda: {agenda.idTipoAgenda}')
    doc.add_paragraph(f'Taller: {agenda.idTaller}')
    doc.add_paragraph(f'Patente Vehículo: {agenda.idVehiculo}')

    # Detalles del Reporte de Pago
    reporte_pago = ReportePago.objects.filter(reserva=agenda).first()
    if reporte_pago:
        doc.add_heading('Detalle del Reporte de Pago', level=2)
        doc.add_paragraph(f'Comentario: {reporte_pago.comentario}')
        doc.add_paragraph(f'Monto a pagar: ${reporte_pago.monto}')

    # Guardar el documento en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Reserva_{agenda.idAgenda}.docx'
    return response

def detalle_reserva(request, id_agenda):
    agenda = get_object_or_404(Agenda, idAgenda=id_agenda)
    reporte_pago = ReportePago.objects.filter(reserva=agenda).first()

    try:
        respuesta = requests.get('https://mindicador.cl/api/')
        monedas = respuesta.json()
        tasa_dolar = Decimal(monedas['dolar']['valor'])
    except requests.exceptions.RequestException as e:
        tasa_dolar = None

    monto_dolares = None
    if reporte_pago and tasa_dolar:
        monto_dolares = round(Decimal(reporte_pago.monto) / tasa_dolar, 2)
        monto_dolares = format(monto_dolares, '.2f')  

    context = {
        'agenda': agenda,
        'reporte_pago': reporte_pago,
        'monto_dolares': monto_dolares,
    }
    return render(request, 'core/detalle_reserva.html', context)

@login_required
def mis_vehiculos(request):
    vehiculos = Vehiculo.objects.filter(idUsuario=request.user)
    return render(request, 'core/mis_vehiculos.html', {'vehiculos': vehiculos})

@login_required
@role_required(["Encargado taller"])
def realizar_ticket(request):
	return render(request, 'core/realizar_ticket.html')

@login_required
def solicitar_taller(request):
	return render(request, 'core/solicitar_taller.html')

@login_required
def register_vehiculo(request):
	return render(request, 'core/register_vehiculo.html')

def register(request):
    form = RegisterForm()
    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)

            user.idRol = rolUsuario.objects.get(nombreRol='Cliente')
            user.idComuna = form.cleaned_data['comuna']
            user.email = user.email
            
            user.save()
            return redirect("index")

    return render(request, 'registration/register.html', {'form': form})

def sobre_nosotros(request):
	return render(request, 'core/sobre_nosotros.html')

@login_required
def talleres(request):
    talleres = Taller.objects.all()
    return render(request, 'core/talleres.html', {'talleres': talleres})

@login_required
@role_required(["Administrador"])
def tickets(request):
    tickets = Ticket.objects.all()
    data = {
        'tickets': tickets
    }
    return render(request, 'core/tickets.html', data)

def get_coordinates(request):
    try:
        error_message=""
        success_message=""
        if request.method == 'POST':
            form = AddressForm(request.POST)
            if form.is_valid():
                address = form.cleaned_data['address']
                api_key = '855d3348e0ee4298b17b00e020377da5'
                base_url = 'https://api.opencagedata.com/geocode/v1/json'
                params = {'q': address, 'key': api_key}
                response = requests.get(base_url, params=params)
                if response.status_code == 200:
                    data = response.json()
                    print(data) 
                    if data['results']:
                        location = data['results'][0]['geometry']
                        address_details = data['results'][0]['components']

                        # Verificar si la dirección está en la Región Metropolitana
                        if 'state' in address_details and address_details['state'] == 'Santiago Metropolitan Region':
                            coordinates = {
                                'lat': location['lat'],
                                'lng': location['lng']
                            }
                            success_message = "La dirección seleccionada es válida"
                            return render(request, 'core/test.html', {'form': form, 'coordinates': coordinates, 'success_message':success_message})
                        else:
                            error_message = "La dirección seleccionada no está en la Región Metropolitana."
                            return render(request, 'core/test.html', {'form': form, 'error_message': error_message})
    except Exception as e:
        print(e)
    form = AddressForm()
    return render(request, 'core/test.html', {'form': form})

def autocomplete_address(request):
    if 'term' in request.GET:
        search_term = request.GET.get('term')
        url = 'https://nominatim.openstreetmap.org/search'
        params = {
            'q': search_term,
            'format': 'json',
            'addressdetails': 1,
            'limit': 5,
        }
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36'
        }
        try:
            response = requests.get(url, params=params, headers=headers)
            response.raise_for_status()  
            suggestions = response.json()
            results = []
            for suggestion in suggestions:
                results.append(suggestion['display_name'])
            return JsonResponse(results, safe=False)
        except requests.RequestException as e:
            print(f"Error during requests to {url}: {str(e)}")
            return JsonResponse([], safe=False)
    return JsonResponse([], safe=False)

@login_required
@role_required(["Administrador"])
def administracion(request):
    return render(request, 'core/administracion.html')

@login_required
def agendar_hora(request, id_taller):
    taller = get_object_or_404(Taller, idTaller=id_taller)

    if request.method == 'POST':
        form = AgendaForm(request.POST, user=request.user, taller=taller)
        if form.is_valid():
            agenda = form.save(commit=False)
            agenda.idTaller = taller
            agenda.cliente = request.user
            agenda.save()
            return redirect('mis_reservas')
    else:
        form = AgendaForm(user=request.user, taller=taller)

    return render(request, 'core/agendar_hora.html', {'form': form, 'taller': taller})

def get_available_hours(request):
    date = request.GET.get('date')
    id_taller = request.GET.get('id_taller')
    taller = get_object_or_404(Taller, idTaller=id_taller)

    if date:
        reserved_hours = Agenda.objects.filter(fechaAtencion=date, idTaller=taller).values_list('horaAtencion', flat=True)
        available_hours = [
            f"{h:02}:00" for h in range(9, 19)
            if time(hour=h) not in reserved_hours
        ]
    else:
        available_hours = [f"{h:02}:00" for h in range(9, 19)]

    return JsonResponse({'available_hours': available_hours})

@login_required
def annadir_vehiculo(request):
    form = VehiculoForm()  

    if request.method == 'POST':
        form = VehiculoForm(request.POST)  

        if form.is_valid():
            vehiculo = form.save(commit=False)
            vehiculo.idUsuario = request.user  
            vehiculo.save()

            return redirect('mis_vehiculos')

    return render(request, 'core/annadir_vehiculo.html', {'form': form})

@login_required
def modificar_vehiculo(request, vehiculo_id):
    vehiculo = Vehiculo.objects.get(idVehiculo=vehiculo_id)

    if request.method == 'POST':
        form = VehiculoForm(request.POST, instance=vehiculo)
        if form.is_valid():
            form.save() 
            return redirect('mis_vehiculos')  
    else:
        form = VehiculoForm(instance=vehiculo)

    return render(request, 'core/modificar_vehiculo.html', {'form': form})

@login_required
def eliminar_vehiculo(request, vehiculo_id):
    vehiculo = Vehiculo.objects.get(idVehiculo = vehiculo_id)
    vehiculo.delete()
    return redirect('mis_vehiculos')
   
@login_required
def administrar_usuarios(request):
    usuarios = UsuarioCustom.objects.all()
    return render(request, 'core/administrar_usuarios.html', {'usuarios': usuarios})

@login_required
def modificar_usuario(request, id):
    usuario = get_object_or_404(UsuarioCustom, id=id)
    if request.method == 'POST':
        form = UsuarioCustomForm(request.POST, instance=usuario)
        if form.is_valid():
            form.save()
            return redirect('administrar_usuarios')
    else:
        form = UsuarioCustomForm(instance=usuario)
    return render(request, 'core/modificar_usuario.html', {'form': form})

@login_required 
@role_required(["Administrador"])
def eliminar_usuario(request, id):
    usuario = UsuarioCustom.objects.get(id = id)
    if request.method == 'POST':
        usuario.delete()
        return redirect('administrar_usuarios')

@login_required 
@role_required(["Administrador"])
def crear_usuario(request):
    if request.method == 'POST':
        form = UsuarioCustomCreationForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('administrar_usuarios')
    else:
        form = UsuarioCustomCreationForm()
    return render(request, 'core/crear_usuario.html', {'form': form})

@login_required
@role_required(["Encargado taller"])
def administrar_mi_taller(request):
    taller_del_usuario = Taller.objects.filter(idUsuario=request.user).first()

    return render(request, 'core/administrar_mi_taller.html', {'taller': taller_del_usuario})

@login_required
@role_required(["Encargado taller"])
def reservas_taller(request, idTaller):
    taller = Taller.objects.get(pk=idTaller)
    reservas = Agenda.objects.filter(idTaller=taller)
    data = {
        'taller': taller,
        'reservas': reservas
    }
    return render(request, 'core/reservas_taller.html', data)

@login_required
def perfil_usuario(request):
    user = request.user
    if request.method == 'POST':
        form = UsuarioCustomPerfilForm(request.POST, instance=user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Perfil actualizado correctamente!')
            return redirect('index')
    else:
        form = UsuarioCustomPerfilForm(instance=user)
    
    return render(request, 'core/perfil_usuario.html', {'form': form})

@role_required(["Encargado taller"])
def generar_reporte_pago(request, idReserva):
    reserva = get_object_or_404(Agenda, pk=idReserva)
    if request.method == 'POST':
        form = ReportePagoForm(request.POST)
        if form.is_valid():
            reporte = form.save(commit=False)
            reporte.reserva = reserva
            reporte.save()

            reserva.estado_id = 2  
            reserva.save()

            return redirect('reservas_taller', idTaller=reserva.idTaller.idTaller)
    else:
        form = ReportePagoForm()
    
    return render(request, 'core/generar_reporte_pago.html', {'form': form, 'reserva': reserva})

@login_required 
@role_required(["Encargado taller"])
def mis_tickets(request):
    user = request.user
    tickets_usuario = Ticket.objects.filter(solicitante=user)
    data = {
        'tickets_usuario':tickets_usuario
    }
    return render(request, 'core/mis_tickets.html', data)

@login_required
@role_required(["Encargado taller"])
def crear_ticket(request):
    if request.method == 'POST':
        form = TicketForm(request.POST)
        if form.is_valid():
            ticket = form.save(commit=False)
            ticket.solicitante = request.user
            estado_pendiente = EstadoTicket.objects.get(NombreEstado='Pendiente')
            ticket.EstadoTicket = estado_pendiente
            ticket.save()
            return redirect('mis_tickets') 
    else:
        form = TicketForm()
    return render(request, 'core/crear_ticket.html', {'form': form})

def aceptar_ticket(request, idTicket):
    ticket = get_object_or_404(Ticket, idTicket=idTicket)
    ticket.EstadoTicket = EstadoTicket.objects.get(NombreEstado='Aceptado')
    ticket.save()
    return redirect('tickets')

def rechazar_ticket(request, idTicket):
    ticket = get_object_or_404(Ticket, idTicket=idTicket)
    ticket.EstadoTicket = EstadoTicket.objects.get(NombreEstado='Rechazado')
    ticket.save()
    return redirect('tickets')

def mensajes(request):
    mensajes = Contacto.objects.all()
    return render(request, 'core/mensajes.html', {'mensajes': mensajes})

def eliminar_mensaje(request, id_mensaje):
    mensaje = get_object_or_404(Contacto, idContacto=id_mensaje)
    mensaje.delete()
    return redirect('mensajes')


